import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

def set_cell_background(cell, fill, color_type=None):
    """
    Sets the background color of a cell.
    """
    shading_elm_1 = OxmlElement('w:shd')
    shading_elm_1.set(qn('w:val'), 'clear')
    shading_elm_1.set(qn('w:color'), 'auto')
    shading_elm_1.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm_1)

def convert_webp_to_png(webp_path, png_path):
    with Image.open(webp_path) as img:
        img.save(png_path, "PNG")

def create_signature_doc():
    doc = Document()
    
    # Paths
    base_dir = r"d:\MindForge"
    logo_webp = os.path.join(base_dir, "public", "logo.webp")
    logo_png = os.path.join(base_dir, "public", "logo.png")
    banner_png = os.path.join(base_dir, "public", "mindforge_banner_930x300.png")
    output_path = os.path.join(base_dir, "mailsignature.docx")

    # Convert Logo
    if os.path.exists(logo_webp):
        convert_webp_to_png(logo_webp, logo_png)
    
    # Create Table
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(0.03)
    table.columns[1].width = Inches(8.0)
    
    # Left Cell - Vertical Bar
    left_cell = table.cell(0, 0)
    set_cell_background(left_cell, "00002B") # Navy

    # Right Cell - Main Content
    right_cell = table.cell(0, 1)
    
    # 1. Logo
    p = right_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run()
    if os.path.exists(logo_png):
        run.add_picture(logo_png, width=Inches(1.8))
        
    # 2. Name
    name_p = right_cell.add_paragraph()
    name_run = name_p.add_run("Balázs Takács")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = "Segoe UI"
    name_p.paragraph_format.space_after = Pt(2)

    # 3. Title
    title_p = right_cell.add_paragraph()
    title_run = title_p.add_run("CEO, Co-Founder & AI Systems Designer")
    title_run.font.size = Pt(11)
    title_run.font.name = "Segoe UI"
    title_p.paragraph_format.space_after = Pt(15)

    # 4. Contact Info
    info_p = right_cell.add_paragraph()
    def add_line(label, text, paragraph):
        run_l = paragraph.add_run(f"{label}: ")
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_l.font.name = "Segoe UI"
        
        run_t = paragraph.add_run(text)
        run_t.font.color.rgb = RGBColor(255, 87, 34) # Accent color
        run_t.font.size = Pt(10)
        run_t.font.name = "Segoe UI"
        paragraph.add_run("\n")

    add_line("T", "+353 85 187 9734", info_p)
    add_line("E", "balazs@mindforgestudio.eu", info_p)
    add_line("W", "www.mindforgestudio.eu", info_p)
    info_p.paragraph_format.space_after = Pt(20)

    # 5. Banner
    banner_p = right_cell.add_paragraph()
    if os.path.exists(banner_png):
        run_b = banner_p.add_run()
        run_b.add_picture(banner_png, width=Inches(6.5))
    banner_p.paragraph_format.space_after = Pt(20)

    # 6. Disclaimer
    disc_p = right_cell.add_paragraph()
    disc_text = (
        "The content of this email is confidential and intended for the recipient specified in message only. "
        "It is strictly forbidden to share any part of this message with any third party, without a written "
        "consent of the sender. If you received this message by mistake, please reply to this message and "
        "follow with its deletion, so that we can ensure such a mistake does not occur in the future."
    )
    disc_run = disc_p.add_run(disc_text)
    disc_run.font.size = Pt(8)
    disc_run.font.name = "Segoe UI"
    disc_run.font.color.rgb = RGBColor(160, 160, 160) # Muted Grey
    disc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")

if __name__ == "__main__":
    create_signature_doc()
